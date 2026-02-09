import streamlit as st
from deep_translator import GoogleTranslator
from docx import Document
from pdf2docx import Converter
from fpdf import FPDF
import io
import os
import time

# --- 1. Konfigurasi Halaman ---
st.set_page_config(page_title="Penerjemah Pro Dua Arah", layout="wide")

st.title("🌐 Penerjemah Dokumen Pro (Dua Arah)")
st.markdown("Support: **DOCX & PDF** | Fitur: **Format & Tabel Tidak Berubah**")
st.write("---")

# --- 2. PENGATURAN BAHASA (Ditaruh di Tengah) ---

# Kita bagi layar jadi 3 kolom: [Kiri, Tengah, Kanan]
# Kita taruh menu di kolom 'Tengah' agar posisinya center
c1, c2, c3 = st.columns([1, 2, 1])

with c2:
    st.info("🛠️ **Pilih Bahasa Terjemahan:**")
    bahasa_opsi = st.selectbox(
        "Mau menerjemahkan dari mana ke mana?",
        ("Indonesia 🇮🇩  ke Inggris 🇬🇧", "Inggris 🇬🇧  ke Indonesia 🇮🇩")
    )

    # Tentukan kode bahasa berdasarkan pilihan
    if bahasa_opsi == "Indonesia 🇮🇩  ke Inggris 🇬🇧":
        SRC_LANG = 'id'
        TGT_LANG = 'en'
    else:
        SRC_LANG = 'en'
        TGT_LANG = 'id'

st.write("---") # Garis pembatas

# --- 3. FUNGSI UTAMA (Backend) ---

def translate_text(text, src, tgt):
    """Menerjemahkan teks dengan penanganan error & retry"""
    if not text or len(text.strip()) < 2:
        return text
    
    # Coba maksimal 3 kali jika gagal
    for attempt in range(3):
        try:
            translator = GoogleTranslator(source=src, target=tgt)
            return translator.translate(text)
        except Exception as e:
            time.sleep(2) # Tunggu 2 detik lalu coba lagi
            continue
            
    return text # Jika gagal total, kembalikan teks asli

def process_docx(file_path_or_buffer, src, tgt):
    """Menerjemahkan isi DOCX tanpa merusak format."""
    doc = Document(file_path_or_buffer)
    
    total_elements = len(doc.paragraphs) + len(doc.tables)
    bar = st.progress(0)
    status = st.empty()
    processed_count = 0
    
    # A. Terjemahkan Paragraf Biasa
    for para in doc.paragraphs:
        if para.text.strip():
            status.text(f"Menerjemahkan paragraf... ({processed_count}/{total_elements})")
            translated = translate_text(para.text, src, tgt)
            para.text = translated
        
        processed_count += 1
        progress = min(processed_count / total_elements, 0.9)
        bar.progress(progress)
            
    # B. Terjemahkan Tabel
    status.text("Sedang menerjemahkan tabel...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        translated = translate_text(paragraph.text, src, tgt)
                        paragraph.text = translated
        processed_count += 1
    
    bar.progress(1.0)
    time.sleep(0.5)
    bar.empty()
    status.success("✅ Selesai! Dokumen siap didownload.")
    return doc

def convert_pdf_to_docx(pdf_file):
    """Mengubah PDF ke DOCX sementara"""
    with open("temp_input.pdf", "wb") as f:
        f.write(pdf_file.getvalue())
    
    cv = Converter("temp_input.pdf")
    cv.convert("temp_output.docx", start=0, end=None)
    cv.close()
    
    return "temp_output.docx"

def save_docx_to_buffer(doc):
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 4. AREA UPLOAD & EKSEKUSI ---

uploaded_file = st.file_uploader("Upload File Anda (PDF atau DOCX)", type=['pdf', 'docx'])

if uploaded_file is not None:
    
    st.write(f"📂 File terdeteksi: **{uploaded_file.name}**")
    
    # Tombol Eksekusi (Dibuat lebar penuh agar gagah)
    if st.button("🚀 Mulai Terjemahkan Sekarang", type="primary", use_container_width=True):
        
        doc_result = None
        
        with st.spinner('Sedang memproses... (Mohon jangan tutup tab ini)'):
            
            # A. PRE-PROCESSING (PDF ke Word)
            if uploaded_file.name.endswith('.pdf'):
                st.info("⚙️ Mengonversi PDF ke format yang bisa diedit...")
                temp_docx = convert_pdf_to_docx(uploaded_file)
                doc_result = process_docx(temp_docx, SRC_LANG, TGT_LANG)
                
                # Bersihkan file sampah
                if os.path.exists("temp_input.pdf"): os.remove("temp_input.pdf")
                if os.path.exists("temp_output.docx"): os.remove("temp_output.docx")
                
            # B. PROSES DOCX LANGSUNG
            elif uploaded_file.name.endswith('.docx'):
                doc_result = process_docx(uploaded_file, SRC_LANG, TGT_LANG)

            # C. OUTPUT DOWNLOAD
            if doc_result:
                docx_buffer = save_docx_to_buffer(doc_result)
                
                full_text = []
                for para in doc_result.paragraphs:
                    full_text.append(para.text)
                txt_output = '\n'.join(full_text)
                
                st.balloons() # Efek balon kalau selesai
                st.success("🎉 Berhasil! Silakan pilih format download:")
                
                d_col1, d_col2 = st.columns(2)
                
                with d_col1:
                    st.download_button(
                        label="📄 Download .DOCX (Layout Asli)",
                        data=docx_buffer,
                        file_name=f"Terjemahan_{uploaded_file.name.split('.')[0]}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    st.caption("✅ Paling Direkomendasikan")

                with d_col2:
                    st.download_button(
                        label="📝 Download .TXT (Teks Saja)",
                        data=txt_output,
                        file_name=f"Terjemahan_{uploaded_file.name.split('.')[0]}.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
