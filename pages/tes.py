import streamlit as st
from deep_translator import GoogleTranslator
from docx import Document
from pdf2docx import Converter
import io
import os
import time

# --- 1. Konfigurasi Halaman ---
st.set_page_config(page_title="Login Penerjemah", layout="wide")

# --- 2. SISTEM LOGIN (Password Protection) ---

def check_password():
    """Mengembalikan True jika user sudah login dengan benar."""
    
    # Cek apakah password sudah disetting di server
    if "PASSWORD" not in st.secrets:
        st.error("⚠️ Password belum diatur di Streamlit Secrets!")
        return False

    def password_entered():
        """Cek input user vs password rahasia"""
        if st.session_state["password_input"] == st.secrets["PASSWORD"]:
            st.session_state["password_correct"] = True
            del st.session_state["password_input"]  # Hapus jejak password
        else:
            st.session_state["password_correct"] = False

    # Inisialisasi status login
    if "password_correct" not in st.session_state:
        st.session_state["password_correct"] = False

    # Jika belum login, tampilkan form input
    if not st.session_state["password_correct"]:
        st.markdown(
            """
            <style>
            .stTextInput {width: 300px; margin: auto;}
            .block-container {padding-top: 5rem;}
            </style>
            """, 
            unsafe_allow_html=True
        )
        st.title("🔒 Halaman Terkunci")
        st.text_input(
            "Masukkan Password Akses:", 
            type="password", 
            on_change=password_entered, 
            key="password_input"
        )
        if "password_correct" in st.session_state and st.session_state["password_correct"] == False:
            st.error("❌ Password salah. Silakan coba lagi.")
        
        return False
    
    return True

# --- 3. JIKA LOGIN BERHASIL, JALANKAN APLIKASI ---

if check_password():
    
    # ==========================================
    # KODE APLIKASI UTAMA DIMULAI DI SINI
    # ==========================================
    
    st.title("🌐 Penerjemah Dokumen Pro (Dua Arah)")
    st.markdown("Support: **DOCX & PDF** | Fitur: **Format & Tabel Tidak Berubah**")
    
    # Tombol Logout
    if st.button("Log Out", type="secondary"):
        st.session_state["password_correct"] = False
        st.rerun()

    st.write("---")

    # --- PENGATURAN BAHASA ---
    c1, c2, c3 = st.columns([1, 2, 1])

    with c2:
        st.info("🛠️ **Pilih Bahasa Terjemahan:**")
        bahasa_opsi = st.selectbox(
            "Mau menerjemahkan dari mana ke mana?",
            ("Indonesia 🇮🇩  ke Inggris 🇬🇧", "Inggris 🇬🇧  ke Indonesia 🇮🇩")
        )

        if bahasa_opsi == "Indonesia 🇮🇩  ke Inggris 🇬🇧":
            SRC_LANG = 'id'; TGT_LANG = 'en'
        else:
            SRC_LANG = 'en'; TGT_LANG = 'id'

    st.write("---")

    # --- FUNGSI ---
    def translate_text(text, src, tgt):
        if not text or len(text.strip()) < 2: return text
        for attempt in range(3):
            try:
                translator = GoogleTranslator(source=src, target=tgt)
                return translator.translate(text)
            except:
                time.sleep(2); continue
        return text

    def process_docx(file_path_or_buffer, src, tgt):
        doc = Document(file_path_or_buffer)
        total = len(doc.paragraphs) + len(doc.tables)
        bar = st.progress(0); status = st.empty(); count = 0
        
        for para in doc.paragraphs:
            if para.text.strip():
                status.text(f"Processing paragraph... ({count}/{total})")
                para.text = translate_text(para.text, src, tgt)
            count += 1; bar.progress(min(count/total, 0.9))
                
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            p.text = translate_text(p.text, src, tgt)
            count += 1
        
        bar.progress(1.0); time.sleep(0.5); bar.empty(); status.success("Selesai!")
        return doc

    def convert_pdf_to_docx(pdf_file):
        with open("temp.pdf", "wb") as f: f.write(pdf_file.getvalue())
        cv = Converter("temp.pdf"); cv.convert("temp.docx"); cv.close()
        return "temp.docx"

    def save_docx(doc):
        buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf

    # --- UPLOAD & ACTION ---
    uploaded_file = st.file_uploader("Upload File (PDF/DOCX)", type=['pdf', 'docx'])

    if uploaded_file:
        if st.button("🚀 Mulai Terjemahkan", type="primary", use_container_width=True):
            doc_res = None
            with st.spinner('Sedang memproses...'):
                if uploaded_file.name.endswith('.pdf'):
                    temp = convert_pdf_to_docx(uploaded_file)
                    doc_res = process_docx(temp, SRC_LANG, TGT_LANG)
                    if os.path.exists("temp.pdf"): os.remove("temp.pdf")
                    if os.path.exists("temp.docx"): os.remove("temp.docx")
                elif uploaded_file.name.endswith('.docx'):
                    doc_res = process_docx(uploaded_file, SRC_LANG, TGT_LANG)

                if doc_res:
                    st.balloons()
                    col1, col2 = st.columns(2)
                    with col1:
                        st.download_button("Download .DOCX", save_docx(doc_res), 
                                        file_name=f"Terjemahan_{uploaded_file.name.split('.')[0]}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        use_container_width=True)
